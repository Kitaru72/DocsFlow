from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.shared import Pt

from app.storage import sanitize_filename


OUTPUT_DIR = Path("output")


def create_document_data(student, report_data):
    return {
        "student": student,
        "report": report_data,
    }


def save_docx(report_data, path=None):
    if path is None:
        path = create_docx_path(report_data)

    doc = Document()

    student = report_data["student"]
    report = report_data["report"]

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    run_discipline = p.add_run(report["discipline"])
    run_discipline.font.size = Pt(18)
    run_discipline.bold = True

    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    run_lab_number = p.add_run(f"Лабораторная работа №{report['lab_number']}")
    run_lab_number.bold = True

    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    run_topic = p.add_run(f"Тема: \"{report['topic']}\"")
    run_topic.bold = True

    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = ALIGN.LEFT
    p.add_run(f"Выполнил: {student['full_name']}")

    p = doc.add_paragraph()
    p.alignment = ALIGN.LEFT
    p.add_run(f"Учебная группа: {student['group']} {student['faculty']}")

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    run_city = p.add_run("Нижневартовск, 2026")
    run_city.bold = True

    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    run_task_heading = p.add_run("Задания")
    run_task_heading.font.size = Pt(18)
    run_task_heading.bold = True

    for num, task in enumerate(report["tasks"], start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"Задание {num}. ")
        run.bold = True
        run = p.add_run(f"{task}")
        run.bold = False

    file_path = Path(path)
    try:
        file_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(file_path)
    except OSError:
        raise ValueError(f"Не удалось сохранить файл {file_path}")

    return file_path


def create_docx_path(report_data):
    last_name = report_data["student"]["full_name"].split()[0]
    group = report_data["student"]["group"]
    lab_number = report_data["report"]["lab_number"]

    filename = f"{last_name}_{group}_лаб{lab_number}.docx"
    formatted_filename = sanitize_filename(filename)

    return OUTPUT_DIR / formatted_filename
